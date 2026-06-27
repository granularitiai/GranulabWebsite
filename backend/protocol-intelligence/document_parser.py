from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass
from pathlib import Path

import fitz
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
MAX_UPLOAD_BYTES = 25 * 1024 * 1024
DEFAULT_MAX_CHARS_FOR_LLM = 350_000
CHUNK_SIZE_CHARS = 30_000


class DocumentParsingError(ValueError):
    pass


class UnsupportedFileTypeError(ValueError):
    pass


class EmptyExtractedTextError(ValueError):
    pass


class FileTooLargeError(ValueError):
    pass


class TokenLimitExceededError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedDocument:
    filename: str
    extension: str
    text: str
    chunks: list[str]
    validation: dict
    warnings: list[str]


PROTOCOL_SIGNALS = {
    "protocol title": r"\bprotocol\s+title\b|\btitle\s+of\s+(the\s+)?protocol\b",
    "protocol number": r"\bprotocol\s+(number|no\.?|id)\b",
    "study phase": r"\bphase\s+(i|ii|iii|iv|1|2|3|4|1/2|2/3)\b",
    "objectives": r"\b(primary|secondary|exploratory)\s+objectives?\b|\bstudy\s+objectives?\b",
    "endpoints": r"\b(primary|secondary|exploratory)\s+endpoints?\b|\befficacy\s+endpoints?\b",
    "inclusion criteria": r"\binclusion\s+criteria\b",
    "exclusion criteria": r"\bexclusion\s+criteria\b",
    "schedule of activities": r"\bschedule\s+of\s+(activities|assessments|events)\b",
    "study arms": r"\b(study|treatment)\s+arms?\b|\barm\s+[a-z0-9]\b",
    "interventions": r"\binterventions?\b|\binvestigational\s+(product|drug|medicinal product)\b",
    "safety monitoring": r"\bsafety\s+(monitoring|assessment|assessments)\b|\badverse\s+events?\b",
    "statistical considerations": r"\bstatistical\s+considerations\b|\bsample\s+size\b|\banalysis\s+population\b",
}


def validate_upload(filename: str, file_bytes: bytes) -> str:
    if not filename:
        raise UnsupportedFileTypeError("Missing filename.")

    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError("Only PDF and DOCX clinical trial protocol files are supported.")

    if not file_bytes:
        raise DocumentParsingError("Uploaded file is empty.")

    if len(file_bytes) > MAX_UPLOAD_BYTES:
        raise FileTooLargeError("File is too large. Maximum supported size is 25 MB.")

    return extension


def parse_document(filename: str, file_bytes: bytes) -> ParsedDocument:
    extension = validate_upload(filename, file_bytes)
    if extension == ".pdf":
        text = extract_pdf_text(file_bytes)
    elif extension == ".docx":
        text = extract_docx_text(file_bytes)
    else:
        raise UnsupportedFileTypeError("Only PDF and DOCX clinical trial protocol files are supported.")

    normalized = normalize_text(text)
    if not normalized:
        raise EmptyExtractedTextError(f"{extension.upper().lstrip('.')} has no extractable text.")

    max_chars_for_llm = get_max_chars_for_llm()
    chunks = chunk_text(normalized)
    if len(normalized) > max_chars_for_llm:
        raise TokenLimitExceededError(
            "Extracted protocol text is too large for this MVP extraction limit. "
            f"The current limit is {max_chars_for_llm:,} extracted characters. "
            "Increase MAX_PROTOCOL_CHARS_FOR_LLM or reduce appendices before uploading."
        )

    validation = classify_protocol_document(normalized)
    warnings: list[str] = []
    if not validation["is_valid_supported_source"]:
        warnings.append(
            "The uploaded document does not strongly appear to be a clinical trial protocol. "
            "Extraction will continue with low confidence where possible."
        )

    return ParsedDocument(
        filename=filename,
        extension=extension,
        text=normalized,
        chunks=chunks,
        validation=validation,
        warnings=warnings,
    )


def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
            return "\n".join(page.get_text("text") for page in pdf)
    except Exception as exc:
        raise DocumentParsingError("Could not parse the uploaded PDF.") from exc


def extract_docx_text(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells: list[str] = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        table_cells.append(text)
        return "\n".join(paragraphs + table_cells)
    except Exception as exc:
        raise DocumentParsingError("Could not parse the uploaded DOCX file.") from exc


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE_CHARS) -> list[str]:
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        boundary = text.rfind("\n", start, end)
        if boundary <= start + int(chunk_size * 0.65):
            boundary = end
        chunks.append(text[start:boundary].strip())
        start = boundary
    return [chunk for chunk in chunks if chunk]


def get_max_chars_for_llm() -> int:
    raw_value = os.getenv("MAX_PROTOCOL_CHARS_FOR_LLM")
    if not raw_value:
        return DEFAULT_MAX_CHARS_FOR_LLM

    try:
        parsed = int(raw_value.replace("_", "").replace(",", ""))
    except ValueError:
        return DEFAULT_MAX_CHARS_FOR_LLM

    return max(50_000, parsed)


def classify_protocol_document(text: str) -> dict:
    lowered = text.lower()
    evidence: list[str] = []

    for label, pattern in PROTOCOL_SIGNALS.items():
        if re.search(pattern, lowered, flags=re.IGNORECASE):
            evidence.append(label)

    confidence = min(1.0, len(evidence) / 8)
    appears_to_be_protocol = confidence >= 0.45 and len(evidence) >= 4

    return {
        "is_valid_supported_source": appears_to_be_protocol,
        "source_confidence": round(confidence, 2),
        "confidence": round(confidence, 2),
        "evidence": evidence,
    }
